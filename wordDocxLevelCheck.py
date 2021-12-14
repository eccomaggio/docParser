#!/usr/bin/env python3
# Written using Python 3.8.2 on MacOS
# by Paul Wakelin July 2021

"""
Function:
Reads a .docx file and compares the text inside against the GEPT wordlist & 
outputs a new (simplified) docx with the original text marked up according to the following 
colour-scheme:
blank = very basic word
    0 > red = not found in the list / offlist
    1 > green = elementary
    2 > orange = intermediate
    3 > blue = hi-int

*** Word .docx format:
in a zipped document containing all kinds of information (see def show_doc_parts)
the text is in the xml/document section
text is arranged like this (but often with many extra tags between the levels):
    paragraph > run > text
    w:p > w:r > w:t

A "run" is a stretch of text that is non-default, e.g. bold / coloured / underlined / italics, etc.

Problems with the current program:

1) undirected
parses EVERYTHING in a document (well, that might not be true: i haven't tested it with tables, and i doubt it can see comments). BUT it currently checks text in e.g. hyperlinks. To get round this, I think I'll have to change from minidom to lxml or elementTree (coz I'll need XPath).

2) words with double entries in the GEPT wordlist (e.g. long verb + long adj.): 
only the second entry will be compared


# USEFUL RESOURCES:

https://towardsdatascience.com/how-to-extract-data-from-ms-word-documents-using-python-ed3fbb48c122
https://www.toptal.com/xml/an-informal-introduction-to-docx
https://docs.python.org/3/library/xml.dom.minidom.html
(https://docs.python.org/3/library/xml.etree.elementtree.html#module-xml.etree.ElementTree)

Great for Minidom
https://python-docx.readthedocs.io/en/latest/
https://python-docx.readthedocs.io/en/latest/user/text.html

For dealing with Docx format
https://virantha.com/2013/08/16/reading-and-writing-microsoft-word-docx-files-with-python/
https://www.programmersought.com/article/42597403131/

https://realpython.com/nltk-nlp-python/
http://www.nltk.org/book/
"""

from collections import namedtuple
import os
import zipfile
import re
import xml.dom.minidom
from xml.dom.minidom import Node
import nltk

from docx import Document
from docx.shared import RGBColor

nltk.download("punkt")
# nltk.download("stopwords")
nltk.download("wordnet")
nltk.download("averaged_perceptron_tagger")


## Step 1: Navigate XML & scrape text  ##################################


def scrape_p_elements(nodelist):
    p_elements = nodelist.getElementsByTagName("w:p")
    all_paragraphs = []
    for paragraph in p_elements:
        all_paragraphs.append(scrape_r_elements(paragraph))
    return all_paragraphs


def scrape_r_elements(paragraph_node):
    r_elements = paragraph_node.getElementsByTagName("w:r")
    text = ""
    for i, run in enumerate(r_elements):
        text += scrape_t_elements(run)
    return text


def scrape_t_elements(run_node):
    t_elements = run_node.getElementsByTagName("w:t")
    text = ""
    for i, t in enumerate(t_elements):
        this_node = t.firstChild
        if this_node.nodeType == Node.TEXT_NODE:
            text += this_node.data
    return text


## Step 2: Tag text  ##################################


def process_text(all_paragraphs):
    processed_paragraphs = []
    for i, paragraph in enumerate(all_paragraphs):
        processed_text = tag_text(make_substitutions(paragraph.lower()))
        processed_paragraphs.append(processed_text)
    return processed_paragraphs


def tag_text(text):
    """
    based on: https://realpython.com/nltk-nlp-python/
    N.B. I use " >" to protect whitespace.
    The space is needed to break up resulting ".<"
    which the lemmatizer would fail to recognize &
    separate from the preceding word, resulting in
    the word not being recognized.
    """
    protect_spaces = text.replace(" ", " <")
    tokens = nltk.word_tokenize(protect_spaces)
    tagged = nltk.pos_tag(tokens)
    marked_up = []
    for item, pos in tagged:
        item_normalized = item.lower()
        ## This allows compound words, phrases & contractions to be checked
        include_compounds = item.translate({ord(c): None for c in "-_'"})
        if include_compounds.isalpha():
            lemma_pos = ""
            lemma = ""
            if pos.startswith("V"):
                lemma_pos = "v"
            elif pos.startswith("J"):
                lemma_pos = "a"
            elif pos.startswith("R"):
                lemma_pos = "r"
            else:
                lemma_pos = "n"
            lemma = lemmatizer.lemmatize(item_normalized, pos=lemma_pos)
            marked_up.append([item, assign_GEPT_level(lemma)])
        else:
            ## For non-word items, i.e. numbers, punctuation, etc.
            marked_up.append([item.translate(text.maketrans(reveal_spaces)), 255])
    return marked_up


def assign_GEPT_level(item):
    if item in GEPT_lookup.keys():
        return GEPT_lookup[item].level
    else:
        return 0


## Reading/Writing files ##################################


def create_GEPT_lookup(lookup_file):
    lookup = {}
    with open(lookup_file, "r", encoding="utf-8") as f:
        for line in f:
            entry = line.strip().split(",")
            info = namedtuple("info", "pos level")
            lookup[entry[0].lower()] = info(entry[1], int(entry[2]))
    return lookup


def make_substitutions(paragraph):
    paragraph = paragraph.translate(paragraph.maketrans(r"’‘’“”–", "'''\"\"-"))

    compounds = {
        "double-tenth day": "Double_Tenth_Day",
        "double tenth day": "Double_Tenth_Day",
        "dragon-boat festival": "Dragon-boat_Festival",
        "dragonboat festival": "Dragon-boat_Festival",
        "dragon boat festival": "Dragon-boat_Festival",
        "hong kong": "Hong_Kong",
        "lantern festival": "Lantern_Festival",
        "mother's day": "Mother_s_Day",
        "mothers day": "Mother_s_Day",
        "mothers' day": "Mother_s_Day",
        "new year's day": "New_Year_s_Day",
        "new years day": "New_Year_s_Day",
        "new year's eve": "New_Year_s_Eve",
        "new years eve": "New_Year_s_Eve",
        "new york": "New_York",
        "republic of china": "Republic_of_China",
        "teacher's day": "Teacher_s_Day",
        "teachers day": "Teacher_s_Day",
        "teachers' day": "Teacher_s_Day",
        "valentine's day": "Valentine_s_Day",
        "valentines day": "Valentine_s_Day",
        "valentines' day": "Valentine_s_Day",
        "o'clock": "o_clock",
        "a\.m\.": "A_M",
        # "am": "A_M",
        "ma'am": "ma_am",
        "p\.m\.": "P_M",
        # "pm": "P_M",
        "mrs.": "Mrs",
        # "mrs": "Mrs",
        "mr\.": "Mr",
        # "mr": "Mr",
        "ms\.": "Ms",
    }

    for key in compounds:
        pattern = re.compile(key)
        paragraph = re.sub(pattern, compounds[key], paragraph)
    return paragraph


def write_xml_file(file_name, xml_output):
    out = homeDir + file_name
    with open(out, "w", encoding="utf-8") as f:
        f.write(xml_output)


def output_docx(all_paragraphs, filename):
    """
    Based on:
    https://python-docx.readthedocs.io/en/latest/
    https://python-docx.readthedocs.io/en/latest/user/text.html
    """
    levels = {
        ## offlist: red rgb ff0000
        0: RGBColor(0xFF, 0x00, 0x00),
        ## elementary: green rgb339933
        1: RGBColor(0x33, 0x99, 0x33),
        ## intermediate: blue rgb 0090cb
        # 2: RGBColor(0xCC, 0x66, 0x00),
        2: RGBColor(0x00, 0x90, 0xCB),
        ## high int: purple rgb 662db9
        3: RGBColor(0x66, 0x2D, 0xB9),
        ## super basic words not in list: grey rgb 8e8e8e
        4: RGBColor(
            0x8E,
            0x8E,
            0x8E,
        ),
    }

    bullet = {"style": "List Bullet"}

    out = Document()
    out.add_heading("Marked up text")

    out.add_paragraph("", **bullet).add_run(
        "Basic contractions & auxiliaries"
    ).font.color.rgb = levels[4]
    out.add_paragraph("", **bullet).add_run("Elementary Level").font.color.rgb = levels[
        1
    ]
    out.add_paragraph("", **bullet).add_run(
        "Intermediate Level"
    ).font.color.rgb = levels[2]
    out.add_paragraph("", **bullet).add_run(
        "High Intermediate Level"
    ).font.color.rgb = levels[3]
    out.add_paragraph("", **bullet).add_run("Off list").font.color.rgb = levels[0]

    for paragraph in all_paragraphs:
        if paragraph:
            p = out.add_paragraph("")
            for item in paragraph:
                if item:
                    if item[1] == 255:
                        p.add_run(item[0])
                    else:
                        run = p.add_run(item[0])
                        font = run.font
                        font.color.rgb = levels[item[1]]
    out.save(filename)


#### initialize global variables  ##############################################

## https://www.codegrepper.com/code-examples/python/get+current+file+name+python
homeDir = os.path.dirname(os.path.realpath(__file__)) + "/"
doc_name = "sample.withTables"
doc_suffix = ".docx"
docFile = doc_name + doc_suffix

lemmatizer = nltk.stem.WordNetLemmatizer()
GEPT_lookup = create_GEPT_lookup(homeDir + "GEPTwordlist.csv")

hide_spaces = {" ": "<"}
reveal_spaces = {"<": " "}

zipped_doc = zipfile.ZipFile(homeDir + docFile)
doc_xml = xml.dom.minidom.parseString(zipped_doc.read("word/document.xml"))


#### MAIN program ##############################################

raw_text = scrape_p_elements(doc_xml)
final_text = process_text(raw_text)

output_docx(final_text, homeDir + doc_name + ".markedup" + doc_suffix)

debug = True
for i, para in enumerate(raw_text):
    if debug:
        if i < 12:
            print(f"orig: {para}\n\n>>>>> {final_text[i]}\n\n")
        else:
            break
    else:
        print(para)
